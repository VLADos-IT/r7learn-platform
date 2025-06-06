import docx


def check_paragraphs_count(expected_docx, comparing_docx, differences):
    if len(expected_docx.paragraphs) != len(expected_docx.paragraphs):
        differences.append('Количество абзацев различается')


def check_paragraphs_indents(expected_paragraph, comparing_paragraph, differences, paragraphs_number):
    expected_format = expected_paragraph.paragraph_format
    comparing_format = comparing_paragraph.paragraph_format

    left_indents_same = expected_format.left_indent != comparing_format.left_indent
    right_indents_same = expected_format.right_indent != comparing_format.right_indent

    if left_indents_same or right_indents_same:
        differences.append(
            f'Отступы в {paragraphs_number}-м абзаце различаются')


def are_font_sizes_equal(font_size_1, font_size_2):
    if font_size_1 is None and font_size_2 is None:
        return True

    if font_size_1 is None or font_size_2 is None:
        return False

    # 0.001 - погрешность сравнения размеров
    return abs(font_size_1 - font_size_2) <= 0.001


def are_styles_equal(font_1, font_2):
    font_styles = ['bold', 'italic', 'underline']

    for style_name in font_styles:
        expected_style = getattr(font_1, style_name, None)
        comparing_style = getattr(font_2, style_name, None)

        if expected_style != comparing_style:
            return False

    return True


def compare_runs(expected_runs, comparing_runs, differences, paragraphs_number):
    if len(expected_runs) != len(comparing_runs):
        differences.append(
            f'Содержимое в {paragraphs_number}-м абзаце различается')
        return

    for expected_run, comparing_run in zip(expected_runs, comparing_runs):
        if expected_run.text.strip() != comparing_run.text.strip():
            differences.append(
                f'Текст в {paragraphs_number}-м абзаце различается')

        if expected_run.font.name != comparing_run.font.name:
            differences.append(
                f'Шрифт в {paragraphs_number}-м абзаце различается')

        if not are_font_sizes_equal(expected_run.font.size, comparing_run.font.size):
            differences.append(
                f'Размер текста в {paragraphs_number}-м абзаце различается')

        if not are_styles_equal(expected_run.font, comparing_run.font):
            differences.append(
                f'Форматирование текста в {paragraphs_number}-м абзаце различается')


def check_paragraphs_content(expected_paragraph, comparing_paragraph, differences, paragraphs_number):
    if expected_paragraph.text.strip() != comparing_paragraph.text.strip():
        differences.append(f'Текст в {paragraphs_number}-м абзаце различается')

    if expected_paragraph.alignment != comparing_paragraph.alignment:
        differences.append(
            f'Выравнивание текста в {paragraphs_number}-м абзаце различается')

    check_paragraphs_indents(
        expected_paragraph, comparing_paragraph, differences, paragraphs_number)
    compare_runs(expected_paragraph.runs, comparing_paragraph.runs,
                 differences, paragraphs_number)


def check_paragraphs(expected_docx, comparing_docx, differences):
    for i in range(len(expected_docx.paragraphs)):
        try:
            expected_paragraph = expected_docx.paragraphs[i]
            comparing_paragraph = comparing_docx.paragraphs[i]
        except IndexError:
            continue

        check_paragraphs_content(
            expected_paragraph, comparing_paragraph, differences, i + 1)


def compare_docx(expected_docx_path: str, comparing_docx_path: str):
    expected_docx = docx.Document(expected_docx_path)
    comparing_docx = docx.Document(comparing_docx_path)

    differences = []

    check_paragraphs_count(expected_docx, comparing_docx, differences)
    check_paragraphs(expected_docx, comparing_docx, differences)

    return differences
